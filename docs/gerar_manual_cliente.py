from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def main() -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for i in range(1, 4):
        hs = doc.styles[f"Heading {i}"]
        hs.font.color.rgb = RGBColor(0x91, 0x3D, 0x4C)
        hs.font.name = "Calibri"
        hs._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    def add_bullets(items):
        for item in items:
            doc.add_paragraph(item, style="List Bullet")

    def add_table(headers, rows):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            for paragraph in table.rows[0].cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = str(val)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        doc.add_paragraph()

    for _ in range(2):
        doc.add_paragraph()

    logo_path = Path(__file__).resolve().parent / "assets" / "logo_isms.png"
    if logo_path.is_file():
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.add_run().add_picture(str(logo_path), width=Cm(3.2))
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Instituto Mais Saúde")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x91, 0x3D, 0x4C)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Manual do Cliente")
    run.bold = True
    run.font.size = Pt(22)

    doc.add_paragraph()

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run("Painel de Integração de Notas Fiscais\nTasy → PR")
    run.font.size = Pt(14)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "Documento de funcionalidades, regras de negócio e operação\n"
        "Versão 1.0 — Julho/2026\n"
        "Confidencial — uso interno do cliente"
    ).font.size = Pt(11)

    doc.add_page_break()

    doc.add_heading("Sumário", level=1)
    for s in [
        "1. Objetivo do sistema",
        "2. Visão geral do fluxo",
        "3. Estabelecimentos e códigos do Tasy",
        "4. Perfis de acesso (Administrador e Usuário)",
        "5. Funcionalidades do painel",
        "6. Rotina automática de extração (6 em 6 minutos)",
        "7. Relatório por e-mail (30 em 30 minutos)",
        "8. Status das notas e tentativas",
        "9. Tipos de erro e tratamento",
        "10. Emissão manual (pendentes e nota específica)",
        "11. Reemissão de notas",
        "12. Filtros, ordenação e paginação",
        "13. Dashboard e exportação de relatório",
        "14. Destinatários de e-mail",
        "15. Configurações operacionais (somente Admin)",
        "16. Logs e auditoria de acessos (somente Admin)",
        "17. Ressalvas e observações importantes",
        "18. Glossário rápido",
    ]:
        doc.add_paragraph(s)

    doc.add_page_break()

    doc.add_heading("1. Objetivo do sistema", level=1)
    doc.add_paragraph(
        "O Painel de Integração de Notas Fiscais tem como objetivo automatizar e controlar "
        "o envio de notas fiscais de entrada do sistema Tasy para o sistema PR (estoque/materiais), "
        "com acompanhamento operacional, tratamento de erros, relatório por e-mail e rastreabilidade "
        "do que foi integrado."
    )
    doc.add_paragraph("Em resumo, o sistema permite:")
    add_bullets(
        [
            "Identificar notas elegíveis no Tasy;",
            "Enviar automaticamente (ou manualmente) essas notas ao PR;",
            "Registrar o resultado da integração (sucesso ou erro);",
            "Marcar a nota no Tasy como integrada após sucesso;",
            "Alertar as equipes por e-mail sobre pendências e ocorrências;",
            "Acompanhar indicadores e exportar relatórios para controle gerencial.",
        ]
    )

    doc.add_heading("2. Visão geral do fluxo", level=1)
    doc.add_paragraph("O fluxo de negócio pode ser entendido assim:")
    add_bullets(
        [
            "1) O sistema consulta o Tasy em busca de notas fiscais elegíveis (ainda não integradas);",
            "2) Cada nota passa por validações (estabelecimento, operação, itens, de-para, lote etc.);",
            "3) A nota é enviada ao PR;",
            "4) Em caso de sucesso, o painel registra a nota como enviada e atualiza a data de integração no Tasy;",
            "5) Em caso de falha, a nota fica disponível para nova tentativa automática ou reemissão manual;",
            "6) Periodicamente, um e-mail resume o que aconteceu para cada estabelecimento.",
        ]
    )
    doc.add_paragraph(
        "Importante: a lista do painel mostra o histórico de notas que já entraram no fluxo de integração. "
        "Notas que existem apenas no Tasy e nunca foram emitidas pelo integrador não aparecem nessa lista "
        "até que sejam capturadas (automaticamente ou por emissão específica)."
    )

    doc.add_heading("3. Estabelecimentos e códigos do Tasy", level=1)
    doc.add_paragraph(
        "O painel opera com quatro estabelecimentos. Cada um corresponde a um código de estabelecimento "
        "no Tasy (cd_estabelecimento):"
    )
    add_table(
        ["Estabelecimento no painel", "Código no Tasy"],
        [
            ["Castelo", "8"],
            ["HRAS", "9"],
            ["HRT (Itaituba)", "7"],
            ["Ponta Porã", "16"],
        ],
    )
    doc.add_paragraph(
        "Esses códigos são usados para filtrar as notas corretas no Tasy. O usuário comum fica "
        "vinculado a um único estabelecimento e só enxerga dados da sua unidade. O administrador "
        "pode visualizar e operar todas as unidades."
    )
    doc.add_heading("3.1 Critérios gerais de elegibilidade da nota", level=2)
    doc.add_paragraph(
        "De forma geral, entram no fluxo as notas que atendem, entre outras, às condições abaixo:"
    )
    add_bullets(
        [
            "Tipo de nota de entrada (EN);",
            "Ainda sem data de integração no Tasy (não integradas);",
            "Operações liberadas: 1 e 39;",
            "Itens com operação 33 são desconsiderados;",
            "Data de emissão a partir de 14/05/2024;",
            "Atualização de estoque nos últimos 31 dias (janela móvel a partir de hoje);",
            "Situações 2 e 3 excluídas.",
        ]
    )
    doc.add_paragraph(
        "Observação: esses critérios podem ser ajustados pela equipe responsável conforme "
        "definição do negócio, mas representam a regra vigente na entrega atual."
    )

    doc.add_heading("4. Perfis de acesso (Administrador e Usuário)", level=1)
    doc.add_heading("4.1 Administrador", level=2)
    doc.add_paragraph("O administrador tem visão global e funções de gestão:")
    add_bullets(
        [
            "Acessa todos os estabelecimentos;",
            "Gerencia usuários do painel;",
            "Liga/desliga a rotina automática e o e-mail por unidade;",
            "Consulta logs de processamento e auditoria de acessos;",
            "Gerencia destinatários de e-mail de qualquer unidade;",
            "Visualiza dashboard consolidado ou por estabelecimento;",
            "Dispara relatório de e-mail manualmente (teste/envio imediato).",
        ]
    )
    doc.add_heading("4.2 Usuário de estabelecimento", level=2)
    doc.add_paragraph("O usuário comum fica restrito ao seu estabelecimento:")
    add_bullets(
        [
            "Emite e acompanha notas apenas da sua unidade;",
            "Consulta o dashboard da própria unidade;",
            "Gerencia somente os e-mails destinatários da própria unidade;",
            "Consulta a Ajuda do painel;",
            "Não acessa Configurações, Logs globais, Acessos/IP nem cadastro de usuários de outras unidades.",
        ]
    )

    doc.add_heading("5. Funcionalidades do painel", level=1)
    doc.add_paragraph("Menu principal e propósito de cada área:")
    add_table(
        ["Área", "Quem acessa", "Para que serve"],
        [
            [
                "Emitir Nota",
                "Admin e Usuário",
                "Emitir pendentes/específicas, filtrar histórico, ver detalhes e reemitir",
            ],
            [
                "Dashboard",
                "Admin e Usuário",
                "Indicadores da integração, erros e exportação CSV",
            ],
            [
                "Destinatários",
                "Admin e Usuário",
                "Cadastrar/editar/excluir e-mails do relatório",
            ],
            ["Usuários", "Somente Admin", "Criar e gerenciar logins"],
            [
                "Logs",
                "Somente Admin",
                "Histórico de processamento e retornos de erro",
            ],
            [
                "Acessos",
                "Somente Admin",
                "Auditoria de IP, usuário, ação e data",
            ],
            [
                "Configurações",
                "Somente Admin",
                "Ligar/desligar rotina e e-mail por unidade",
            ],
            [
                "Ajuda",
                "Admin e Usuário",
                "Orientações de uso (conteúdo adaptado ao perfil)",
            ],
        ],
    )

    doc.add_heading("6. Rotina automática de extração (6 em 6 minutos)", level=1)
    doc.add_paragraph(
        "Existe uma rotina automática que, a cada 6 minutos, busca no Tasy as notas elegíveis "
        "dos estabelecimentos habilitados e as encaminha para integração com o PR."
    )
    doc.add_heading("6.1 Como a rotina é controlada", level=2)
    add_bullets(
        [
            "Controle geral (liga/desliga de toda a rotina): responsabilidade da operação;",
            "Controle por estabelecimento: tela Configurações (somente Admin), coluna Scheduler (6 min);",
            "Para uma unidade processar automaticamente, os dois controles precisam estar ligados "
            "(geral + da unidade);",
            "Se o controle geral estiver ligado, mas todas as unidades estiverem desligadas, nenhuma nota "
            "será capturada automaticamente.",
        ]
    )
    doc.add_heading("6.2 O que a rotina faz a cada ciclo", level=2)
    add_bullets(
        [
            "Consulta notas elegíveis no Tasy da unidade;",
            "Encaminha essas notas para processamento/envio ao PR;",
            "Não depende de alguém estar logado no painel;",
            "Não substitui a emissão manual (pendentes ou específica), que continua disponível "
            "mesmo com a rotina desligada.",
        ]
    )

    doc.add_heading("7. Relatório por e-mail (30 em 30 minutos)", level=1)
    doc.add_paragraph(
        "A cada 30 minutos, o sistema pode enviar um relatório por e-mail por estabelecimento, "
        "desde que o envio esteja habilitado para aquela unidade e existam destinatários cadastrados."
    )
    doc.add_heading("7.1 Controles do e-mail", level=2)
    add_bullets(
        [
            "Controle geral do envio automático: responsabilidade da operação;",
            "Controle por unidade: Configurações → Relatório e-mail (30 min) — somente Admin;",
            "Destinatários: aba Destinatários (Admin vê todas as unidades; Usuário só a sua);",
            "Envio manual imediato: botão em Configurações (“Enviar relatório agora”) — somente Admin;",
            "Se não houver destinatários, o e-mail não é enviado;",
            "No modo automático, se não houver ocorrência nova relevante, o e-mail não é disparado.",
        ]
    )
    doc.add_heading("7.2 Seções do e-mail", level=2)
    add_table(
        ["Seção", "Comportamento"],
        [
            ["Notas integradas com sucesso", "Entra no e-mail uma única vez"],
            [
                "Notas não integradas (sem bloqueio de de-para/lote)",
                "Entra no e-mail uma única vez",
            ],
            ["Itens sem de-para", "Pode repetir a cada ciclo até ser resolvido"],
            [
                "Itens com necessidade de lote",
                "Pode repetir a cada ciclo até ser resolvido",
            ],
            ["Erros de retorno do PR", "Entra no e-mail uma única vez"],
        ],
    )
    doc.add_paragraph(
        "Para evitar que a mesma nota apareça em várias seções ao mesmo tempo, há prioridade: "
        "sem de-para → sem lote → não integrada."
    )
    doc.add_paragraph(
        "Isso reduz ruído: o time continua sendo lembrado das pendências operacionais (de-para/lote), "
        "mas não recebe repetidamente o mesmo aviso de nota já integrada ou já retornada pelo PR."
    )

    doc.add_heading("8. Status das notas e tentativas", level=1)
    add_table(
        ["Status no painel", "Significado"],
        [
            [
                "Pendente (pending)",
                "Nota entrou no fluxo e ainda não concluiu o processamento final",
            ],
            ["Enviado (sent)", "Integração concluída com sucesso no PR"],
            [
                "Aguardando retry (retry_pending)",
                "Falhou e aguarda nova tentativa automática",
            ],
            [
                "Falha definitiva (dead_letter)",
                "Esgotou as tentativas automáticas; exige ação (reemitir após correção)",
            ],
        ],
    )
    doc.add_paragraph(
        "Em caso de erro no envio ao PR, o sistema realiza até 3 tentativas automáticas. "
        "Se todas falharem, a nota fica em falha definitiva e pode ser reemitida manualmente "
        "depois que a causa for corrigida."
    )

    doc.add_heading("9. Tipos de erro e tratamento", level=1)
    add_table(
        ["Tipo de erro", "O que significa", "O que fazer"],
        [
            [
                "Sem de-para",
                "Material do Tasy sem vínculo cadastrado no PR",
                "Cadastrar o de-para no PR e usar Reemitir",
            ],
            [
                "Sem lote",
                "Item exige lote e a informação não veio no Tasy",
                "Corrigir o lote no Tasy e reemitir",
            ],
            [
                "Retorno PR",
                "O PR recusou a nota (ex.: já integrada, validação)",
                "Analisar a mensagem de retorno; corrigir a causa se necessário",
            ],
            [
                "Outro",
                "Erro não classificado nas categorias acima",
                "Analisar o detalhe do erro no painel/logs",
            ],
        ],
    )
    doc.add_paragraph(
        "Após sucesso no PR, o sistema tenta atualizar no Tasy o campo de data de integração "
        "(dt_integracao). Com isso, a nota deixa de ser elegível para nova captura automática. "
        "Se a atualização dessa data falhar, a nota permanece marcada como enviada no painel, "
        "com alerta na mensagem de retorno, para não perder o histórico de sucesso."
    )

    doc.add_heading("10. Emissão manual (pendentes e nota específica)", level=1)
    doc.add_heading("10.1 Emitir pendentes", level=2)
    doc.add_paragraph(
        "Busca no Tasy todas as notas elegíveis do estabelecimento selecionado e as envia para o fluxo "
        "de integração. Não reenvia notas que já constam no histórico de processamento do painel. "
        "Funciona mesmo com a rotina automática desligada."
    )
    doc.add_heading("10.2 Emitir nota específica", level=2)
    doc.add_paragraph(
        "Permite informar o NR Sequência da nota no Tasy, consultar se ela é válida/elegível e, "
        "se estiver correta, confirmar a emissão. Útil para reprocessar um caso pontual ou validar "
        "uma nota específica antes do envio."
    )

    doc.add_heading("11. Reemissão de notas", level=1)
    doc.add_paragraph("O botão Reemitir aparece quando a nota:")
    add_bullets(
        [
            "Está em Aguardando retry ou Falha definitiva; e",
            "Possui NR Sequência registrado.",
        ]
    )
    doc.add_paragraph(
        "A reemissão busca novamente os dados no Tasy e recoloca a nota no fluxo. "
        "Notas já enviadas com sucesso (sent) não devem ser reemitidas por esse botão. "
        "Se a nota não tiver NR Sequência no histórico, utilize Emitir nota específica."
    )

    doc.add_heading("12. Filtros, ordenação e paginação", level=1)
    doc.add_paragraph("Na tela Emitir Nota:")
    add_bullets(
        [
            "A lista não carrega automaticamente ao abrir a página — é necessário aplicar filtros;",
            "Filtros disponíveis: NF, NR Sequência, Fornecedor, Status, Tipo de erro, Data NF (de/até);",
            "Ordenação: por NR Sequência (maior → menor) ou por Data NF (mais recente);",
            "Paginação: resultados em páginas (padrão 50 registros);",
            "Limpar filtros esvazia a lista até uma nova pesquisa;",
            "Clique em uma linha abre o detalhe da nota (itens, lotes, status de de-para e totais).",
        ]
    )

    doc.add_heading("13. Dashboard e exportação de relatório", level=1)
    doc.add_paragraph("O Dashboard apresenta uma visão gerencial da integração:")
    add_bullets(
        [
            "Total de notas no período;",
            "Quantidade integradas com sucesso e percentual de sucesso;",
            "Retry pendente e falha definitiva;",
            "Distribuição por status e por tipo de erro;",
            "Visão por estabelecimento (Admin);",
            "Evolução diária (integradas vs erros);",
            "Lista dos erros mais recentes.",
        ]
    )
    doc.add_paragraph(
        "Também há a opção Exportar CSV, que gera um arquivo para Excel com os principais campos "
        "de controle da nota (estabelecimento, NF, sequência, fornecedor, datas, status, tentativas, "
        "tipo de erro, mensagem de erro, retorno do PR, datas de criação/atualização no painel). "
        "O usuário comum exporta apenas a sua unidade; o Admin pode filtrar por unidade ou exportar "
        "o conjunto filtrado."
    )

    doc.add_heading("14. Destinatários de e-mail", level=1)
    doc.add_paragraph(
        "A aba Destinatários é separada das Configurações justamente para que os estabelecimentos "
        "possam manter a lista de e-mails sem alterar ligar/desligar da rotina automática."
    )
    add_bullets(
        [
            "É possível adicionar, editar e excluir e-mails;",
            "Usuário comum só gerencia os e-mails do próprio estabelecimento;",
            "Admin escolhe a unidade e gerencia qualquer lista;",
            "Sem destinatários cadastrados, o relatório daquela unidade não é enviado;",
            "O ligar/desligar do disparo automático continua exclusivo do Admin (Configurações).",
        ]
    )

    doc.add_heading("15. Configurações operacionais (somente Admin)", level=1)
    doc.add_paragraph("Na tela Configurações, o Admin controla por estabelecimento:")
    add_bullets(
        [
            "Scheduler (6 min): liga/desliga a captura automática da unidade;",
            "Relatório e-mail (30 min): liga/desliga o envio automático de e-mail da unidade;",
            "Envio imediato de relatório (teste ou demanda pontual).",
        ]
    )
    doc.add_paragraph(
        "A emissão manual pelo painel (Emitir pendentes / Emitir nota específica) não depende "
        "desses interruptores."
    )

    doc.add_heading("16. Logs e auditoria de acessos (somente Admin)", level=1)
    doc.add_heading("16.1 Logs de processamento", level=2)
    doc.add_paragraph(
        "Exibem o histórico operacional das notas (status, tipo de erro, tentativas e retorno do PR), "
        "com filtros e paginação. Serve para acompanhar falhas e entender o que ocorreu em cada tentativa."
    )
    doc.add_heading("16.2 Acessos", level=2)
    doc.add_paragraph(
        "Registra auditoria de uso do painel: data/hora, usuário, IP, ação realizada, caminho acessado "
        "e status da operação. Há filtros por data, usuário, IP, perfil, estabelecimento e ação, "
        "além de um resumo de IPs únicos para visão consolidada — sem necessidade de analisar "
        "registro por registro."
    )

    doc.add_heading("17. Ressalvas e observações importantes", level=1)
    add_bullets(
        [
            "Cadastro de de-para no PR é pré-requisito para integração de materiais; sem vínculo, a nota falha;",
            "Itens que exigem lote precisam ter lote informado no Tasy;",
            "Nota já integrada no Tasy (com data de integração preenchida) não volta para a captura automática;",
            "Nota já existente no PR pode retornar erro de “já integrada” / lançamento duplicado;",
            "Somente operações liberadas (1 e 39) e demais filtros do perfil entram no fluxo;",
            "O painel não é um espelho completo do Tasy: mostra o que já passou pelo integrador;",
            "E-mails de pendência (sem de-para / sem lote) podem se repetir até a correção; sucessos e retornos PR entram uma vez;",
            "Reemitir só deve ser usado depois de corrigida a causa do erro;",
            "Alterações de destinatários no painel passam a valer para os próximos disparos de e-mail;",
            "O acesso ao painel deve ser feito com usuário individual, pois há auditoria de IP e ações;",
            "Em ambiente publicado por endereço/DDNS, utilize sempre a URL oficial fornecida pela operação.",
        ]
    )

    doc.add_heading("18. Glossário rápido", level=1)
    add_table(
        ["Termo", "Significado"],
        [
            ["Tasy", "Sistema de origem das notas fiscais"],
            ["PR", "Sistema de destino da integração (estoque/materiais)"],
            ["NR Sequência", "Identificador sequencial da nota no Tasy"],
            ["De-para", "Vínculo entre código de material do Tasy e do PR"],
            [
                "dt_integracao",
                "Data/hora em que a nota foi marcada como integrada no Tasy",
            ],
            ["Emitir pendentes", "Captura em lote das notas elegíveis da unidade"],
            ["Emitir nota específica", "Captura pontual por NR Sequência"],
            ["Reemitir", "Nova tentativa manual após falha"],
            [
                "Dead letter / Falha definitiva",
                "Nota que esgotou as tentativas automáticas",
            ],
            ["Scheduler", "Rotina automática de captura a cada 6 minutos"],
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "— Fim do documento —\n"
        "Em caso de dúvidas operacionais, utilize a aba Ajuda no painel ou contate o suporte "
        "responsável pela integração."
    )
    run.italic = True
    run.font.size = Pt(10)

    out_dir = Path(__file__).resolve().parent
    path = out_dir / "Manual_Cliente_Painel_Integracao_Notas_Fiscais.docx"
    doc.save(path)
    print(path)


if __name__ == "__main__":
    main()
