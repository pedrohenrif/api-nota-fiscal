"""Gera docs/Documentacao_Tecnica_Integracao_Notas_Fiscais.docx a partir de DOCUMENTACAO_TECNICA.md."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "DOCUMENTACAO_TECNICA.md"
OUT_PATH = Path(__file__).resolve().parent / "Documentacao_Tecnica_Integracao_Notas_Fiscais.docx"
LOGO_PATH = Path(__file__).resolve().parent / "assets" / "logo_isms.png"
BRAND = RGBColor(0x91, 0x3D, 0x4C)


def _strip_md(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("**", "").replace("`", "")
    return text.strip()


def _setup_styles(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for i in range(1, 4):
        hs = doc.styles[f"Heading {i}"]
        hs.font.color.rgb = BRAND
        hs.font.name = "Calibri"
        hs._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def _add_runs_with_bold(paragraph, text: str) -> None:
    """Suporta **negrito** simples no meio do texto."""
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            # links markdown → só o label
            part = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", part)
            paragraph.add_run(part)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = _strip_md(h)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        padded = list(row) + [""] * max(0, cols - len(row))
        for c_idx in range(cols):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = _strip_md(padded[c_idx] if c_idx < len(padded) else "")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def _add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)


def _add_cover(doc: Document) -> None:
    for _ in range(2):
        doc.add_paragraph()

    if LOGO_PATH.is_file():
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.add_run().add_picture(str(LOGO_PATH), width=Cm(3.2))
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Instituto Mais Saúde")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = BRAND

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Documentação Técnica")
    run.bold = True
    run.font.size = Pt(22)

    doc.add_paragraph()

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run("Integração de Notas Fiscais\nTasy → PR")
    run.font.size = Pt(14)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "Documento para desenvolvedores e operação técnica\n"
        "Versão 1.0 — Julho/2026\n"
        "Fonte: DOCUMENTACAO_TECNICA.md"
    ).font.size = Pt(11)

    doc.add_page_break()


def convert_md_to_docx(md_text: str, doc: Document) -> None:
    lines = md_text.splitlines()
    i = 0
    # Pula o H1 do markdown (capa já cobre)
    if lines and lines[0].startswith("# "):
        i = 1
        while i < len(lines) and lines[i].strip() in ("", "---"):
            i += 1
        # pula bloco intro até o primeiro ## se a capa já descreveu
        # Mantém o intro como texto após sumário implícito

    in_code = False
    code_lines: list[str] = []
    table_buf: list[str] = []

    def flush_table() -> None:
        nonlocal table_buf
        if not table_buf:
            return
        parsed: list[list[str]] = []
        for raw in table_buf:
            cells = [c.strip() for c in raw.strip().strip("|").split("|")]
            if all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
                continue
            parsed.append(cells)
        table_buf = []
        if len(parsed) >= 2:
            _add_table(doc, parsed[0], parsed[1:])
        elif parsed:
            p = doc.add_paragraph(_strip_md(" | ".join(parsed[0])))

    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            flush_table()
            if in_code:
                _add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith("|"):
            table_buf.append(line)
            i += 1
            continue
        else:
            flush_table()

        if line.strip() == "---":
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(_strip_md(line[4:]), level=3)
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(_strip_md(line[3:]), level=2)
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(_strip_md(line[2:]), level=1)
            i += 1
            continue

        if re.match(r"^\d+\.\s+", line.strip()):
            text = re.sub(r"^\d+\.\s+", "", line.strip())
            p = doc.add_paragraph(style="List Number")
            _add_runs_with_bold(p, text)
            i += 1
            continue

        if line.strip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_bold(p, line.strip()[2:])
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.strip().startswith("*") and line.strip().endswith("*") and not line.strip().startswith("**"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(_strip_md(line.strip().strip("*")))
            run.italic = True
            run.font.size = Pt(10)
            i += 1
            continue

        p = doc.add_paragraph()
        _add_runs_with_bold(p, line)
        i += 1

    flush_table()
    if in_code and code_lines:
        _add_code_block(doc, code_lines)


def main() -> None:
    if not MD_PATH.is_file():
        raise SystemExit(f"Markdown não encontrado: {MD_PATH}")

    doc = Document()
    _setup_styles(doc)
    _add_cover(doc)
    convert_md_to_docx(MD_PATH.read_text(encoding="utf-8"), doc)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
