"""Gera docs/Documentacao_Projeto_Integracao_Notas_Fiscais.docx

Documento para o cliente: regras de negocio, tratativas e painel — sem codigo.
Capa: logo ISMS (assets/logo_isms.png). Logo GHR: docs/assets/logo_ghr.png (quando disponivel).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BRAND = RGBColor(0x91, 0x3D, 0x4C)
OUT = Path(__file__).resolve().parent / "Documentacao_Projeto_Integracao_Notas_Fiscais.docx"
ASSETS = Path(__file__).resolve().parent / "assets"
LOGO_ISMS = ASSETS / "logo_isms.png"
LOGO_GHR_CANDIDATES = (
    ASSETS / "logo_ghr.png",
    ASSETS / "logo_ghr.jpg",
    ASSETS / "logo_ghr.jpeg",
    ASSETS / "GHR-Logo.jpeg",
    ASSETS / "GHR-Logo.jpg",
    ASSETS / "GHR-Logo.png",
)


def _resolve_logo_ghr() -> Path | None:
    for path in LOGO_GHR_CANDIDATES:
        if path.is_file():
            return path
    return None


def main() -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for i in range(1, 4):
        hs = doc.styles[f"Heading {i}"]
        hs.font.color.rgb = BRAND
        hs.font.name = "Calibri"
        hs._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    def add_bullets(items: list[str]) -> None:
        for item in items:
            doc.add_paragraph(item, style="List Bullet")

    def add_table(headers: list[str], rows: list[list[str]]) -> None:
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            for paragraph in table.rows[0].cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = str(val)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        doc.add_paragraph()

    # --- Capa ---
    for _ in range(2):
        doc.add_paragraph()

    logos = doc.add_paragraph()
    logos.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_ISMS.is_file():
        logos.add_run().add_picture(str(LOGO_ISMS), width=Cm(3.0))
    logo_ghr = _resolve_logo_ghr()
    if logo_ghr is not None:
        logos.add_run("    ")
        logos.add_run().add_picture(str(logo_ghr), width=Cm(3.8))
    else:
        note = doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = note.add_run("[Espaço reservado para a logo GHR Tech]")
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Instituto Mais Saúde")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = BRAND

    ghr = doc.add_paragraph()
    ghr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ghr.add_run("GHR Tech")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Documentação do Projeto")
    run.bold = True
    run.font.size = Pt(22)

    doc.add_paragraph()

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run(
        "Integração de Notas Fiscais\nTasy → PR\n"
        "Regras de negócio, tratativas e painel operacional"
    )
    run.font.size = Pt(14)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "Documento destinado ao cliente (sem detalhamento técnico de código)\n"
        "Versão 1.1 — Agosto/2026\n"
        "Confidencial"
    ).font.size = Pt(11)

    doc.add_page_break()

    doc.add_heading("Sumário", level=1)
    for s in [
        "1. Apresentação e objetivo",
        "2. Escopo e estabelecimentos",
        "3. Visão geral do fluxo",
        "4. Regras de negócio (elegibilidade)",
        "5. Tratativas de integração (de-para, lote, valores, write-back)",
        "6. Status das notas e tipos de erro",
        "7. O painel web",
        "8. Rotinas automáticas (extração e e-mail)",
        "9. Relatório por e-mail",
        "10. Perfis de acesso",
        "11. Ressalvas e pontos de atenção",
        "12. Glossário",
    ]:
        doc.add_paragraph(s)

    doc.add_page_break()

    # 1
    doc.add_heading("1. Apresentação e objetivo", level=1)
    doc.add_paragraph(
        "Este documento descreve o projeto de Integração de Notas Fiscais desenvolvido pela "
        "GHR Tech para o Instituto Mais Saúde. O foco é o entendimento de negócio: o que o "
        "sistema faz, quais regras aplica, como trata falhas e como o painel é utilizado no dia a dia."
    )
    doc.add_paragraph(
        "O objetivo é automatizar e controlar o envio de notas fiscais de entrada do sistema "
        "Tasy para o sistema PR (estoque/materiais), com acompanhamento operacional, tratamento "
        "de erros, alertas por e-mail e rastreabilidade do que foi integrado."
    )
    doc.add_paragraph("Em resumo, a solução permite:")
    add_bullets(
        [
            "Identificar notas elegíveis no Tasy;",
            "Enviar automaticamente ou manualmente essas notas ao PR;",
            "Registrar sucesso ou erro de cada tentativa;",
            "Marcar a nota no Tasy como integrada após sucesso no PR;",
            "Alertar as equipes por e-mail sobre pendências e ocorrências;",
            "Acompanhar indicadores e exportar relatórios no painel.",
        ]
    )

    # 2
    doc.add_heading("2. Escopo e estabelecimentos", level=1)
    doc.add_paragraph(
        "A integração contempla quatro estabelecimentos, cada um com código próprio no Tasy:"
    )
    add_table(
        ["Estabelecimento", "Código no Tasy"],
        [
            ["Castelo", "8"],
            ["HRAS", "9"],
            ["HRT (Itaituba)", "7"],
            ["Ponta Porã", "16"],
        ],
    )
    doc.add_paragraph(
        "Cada unidade pode ter a rotina automática e o e-mail ligados ou desligados de forma "
        "independente. Usuários comuns enxergam apenas o estabelecimento ao qual estão vinculados; "
        "administradores têm visão de todas as unidades."
    )

    # 3
    doc.add_heading("3. Visão geral do fluxo", level=1)
    add_bullets(
        [
            "1) O sistema consulta o Tasy em busca de notas elegíveis (ainda sem data de integração);",
            "2) Aplica as regras de negócio (tipo, operação, datas, itens, local de estoque etc.);",
            "3) Valida de-para dos materiais e informações de lote quando aplicável;",
            "4) Envia a nota ao PR;",
            "5) Em sucesso: registra no painel como enviada e grava a data de integração no Tasy;",
            "6) Em falha: classifica o tipo de erro, tenta novamente (até o limite) e permite reemissão manual;",
            "7) Periodicamente, um e-mail resume pendências e erros por estabelecimento.",
        ]
    )
    doc.add_paragraph(
        "O painel mostra o histórico das notas que já entraram no fluxo do integrador. "
        "Notas que existem só no Tasy e nunca foram capturadas não aparecem na lista até serem "
        "emitidas (automática ou manualmente)."
    )

    # 4
    doc.add_heading("4. Regras de negócio (elegibilidade)", level=1)
    doc.add_paragraph(
        "Entram no fluxo, em regra, as notas que atendem às condições abaixo (regra vigente "
        "na entrega atual):"
    )
    add_bullets(
        [
            "Tipo de nota de entrada (EN);",
            "Sem data de integração no Tasy (ainda não integradas);",
            "Operações liberadas: 1 e 39;",
            "Itens vinculados à operação 33 são desconsiderados;",
            "Situações 2 e 3 excluídas;",
            "Data de emissão a partir de 14/05/2024;",
            "Data de atualização de estoque a partir de 05/08/2026 (piso operacional);",
            "Itens com local de estoque 104 não são elegíveis para integração;",
            "A nota precisa ter ao menos um item elegível após os filtros acima.",
        ]
    )
    doc.add_paragraph(
        "Na consulta de nota específica no painel, quando não há itens elegíveis, o sistema "
        "exibe o diagnóstico dos locais de estoque encontrados (por exemplo, todos em 104), "
        "evitando a necessidade de consultar o banco manualmente."
    )

    # 5
    doc.add_heading("5. Tratativas de integração", level=1)

    doc.add_heading("5.1 De-para de materiais", level=2)
    doc.add_paragraph(
        "Antes do envio definitivo, o sistema consulta o PR para verificar se o código do "
        "material do Tasy possui vínculo (de-para) cadastrado. O código apresentado no e-mail "
        "e no painel é o código do material no Tasy (sede), não o código interno do PR."
    )
    add_bullets(
        [
            "Com vínculo: a nota segue no fluxo;",
            "Sem vínculo: a nota é classificada como “sem de-para” e permanece pendente até o cadastro no PR;",
            "O de-para deve existir no ambiente e unidade corretos (produção × homologação; token da unidade).",
        ]
    )

    doc.add_heading("5.2 Lote", level=2)
    doc.add_paragraph(
        "Itens que exigem lote precisam ter lote informado no Tasy. Notas/itens sem lote "
        "elegível são classificados como “sem lote” e podem repetir no e-mail até a correção."
    )

    doc.add_heading("5.3 Valores e desconto", level=2)
    doc.add_paragraph(
        "No envio ao PR, a solução utiliza o total da nota, o desconto do cabeçalho e os "
        "valores dos itens conforme registrados no Tasy. O PR valida se o valor total da NF "
        "é coerente com a soma dos valores dos produtos."
    )
    doc.add_paragraph(
        "Quando há desconto relevante no cabeçalho e os itens permanecem com valores brutos, "
        "o PR pode rejeitar a nota com mensagem do tipo “valor de entrada divergente "
        "(Estoque.NF <> Estoque.ProdutoNF)”. Nesse caso a divergência existe nos dados da nota "
        "no Tasy em relação à regra de fechamento do PR — não é um falso positivo do painel. "
        "A tratativa operacional é ajustar os valores no Tasy (ou alinhar a regra com o PR) e reemitir."
    )

    doc.add_heading("5.4 Marcação no Tasy após sucesso (write-back)", level=2)
    doc.add_paragraph(
        "Após integração bem-sucedida no PR, o sistema grava a data/hora de integração na nota "
        "no Tasy. Isso evita que a mesma nota volte a ser capturada automaticamente. No detalhe "
        "da nota no painel, status “sent” é apresentado como integrada no PR (indicação visual "
        "de sucesso), mantendo os dados da nota visíveis mesmo com a data de integração já preenchida."
    )

    doc.add_heading("5.5 Tentativas e reemissão", level=2)
    add_bullets(
        [
            "Falhas transitórias ou de retorno do PR podem gerar novas tentativas automáticas até o limite configurado;",
            "Após esgotar tentativas, a nota fica em falha definitiva (dead letter) e pode ser reemitida manualmente no painel após correção da causa;",
            "Reemitir só deve ser usado depois de corrigida a pendência (de-para, lote, valor etc.).",
        ]
    )

    # 6
    doc.add_heading("6. Status das notas e tipos de erro", level=1)
    add_table(
        ["Status no painel", "Significado"],
        [
            ["pending", "Aguardando / em processamento inicial"],
            ["sent", "Integrada com sucesso no PR"],
            ["retry_pending", "Falha com nova tentativa prevista"],
            ["dead_letter", "Esgotou tentativas; exige ação / reemissão"],
        ],
    )
    add_table(
        ["Tipo de erro", "Tratativa típica"],
        [
            [
                "sem_depara",
                "Cadastrar vínculo do material Tasy ↔ PR e reemitir; pode repetir no e-mail até corrigir",
            ],
            [
                "sem_lote",
                "Informar lote no Tasy e reemitir; pode repetir no e-mail até corrigir",
            ],
            [
                "retorno_pr",
                "Analisar mensagem do PR (ex.: valor divergente, nota já existente); corrigir causa e reemitir",
            ],
            ["outro", "Analisar mensagem e logs; tratar conforme o caso"],
        ],
    )

    # 7
    doc.add_heading("7. O painel web", level=1)
    doc.add_paragraph(
        "O painel é a interface operacional da integração. Principais áreas:"
    )
    add_table(
        ["Área", "Para que serve"],
        [
            [
                "Emitir Nota",
                "Emitir pendentes ou nota específica, filtrar histórico, abrir detalhe (itens, de-para, lotes, totais) e reemitir",
            ],
            [
                "Dashboard",
                "Indicadores de integração, erros e exportação CSV",
            ],
            [
                "Destinatários",
                "Cadastro dos e-mails que recebem o relatório da unidade",
            ],
            [
                "Configurações (Admin)",
                "Ligar/desligar scheduler e e-mail por unidade; intervalo do e-mail (6 ou 30 min)",
            ],
            [
                "Usuários (Admin)",
                "Gestão de logins e vínculos a estabelecimentos",
            ],
            [
                "Logs / Acessos (Admin)",
                "Histórico de processamento e auditoria de uso (incluindo IP)",
            ],
            ["Ajuda", "Orientação de uso conforme o perfil"],
        ],
    )
    doc.add_paragraph(
        "No detalhe da nota é possível conferir identificação, cabeçalho, itens com status de "
        "de-para, lotes e totais (incluindo comparação visual entre soma dos itens e valor total "
        "da NF). Notas já integradas no PR exibem indicação de sucesso e mantêm os dados abertos "
        "para consulta."
    )

    # 8
    doc.add_heading("8. Rotinas automáticas", level=1)
    doc.add_heading("8.1 Extração automática", level=2)
    add_bullets(
        [
            "Ciclo padrão a cada 6 minutos;",
            "Existe um interruptor geral e um interruptor por estabelecimento no painel;",
            "Só entram no ciclo as unidades com scheduler ligado;",
            "A emissão manual (pendentes ou específica) funciona mesmo com o scheduler desligado.",
        ]
    )
    doc.add_heading("8.2 Relatório por e-mail", level=2)
    add_bullets(
        [
            "Intervalo configurável no painel: 6 ou 30 minutos;",
            "Também depende do interruptor por estabelecimento;",
            "Disparo automático ocorre quando há pendências/erros (sem de-para, sem lote, retorno PR, não integradas);",
            "Notas já emitidas com sucesso não disparam o ciclo sozinhas;",
            "Horário do e-mail segue o fuso de Brasília.",
        ]
    )

    # 9
    doc.add_heading("9. Relatório por e-mail", level=1)
    doc.add_paragraph("O e-mail consolida, por estabelecimento, seções como:")
    add_bullets(
        [
            "Itens / notas sem de-para;",
            "Sem lote;",
            "Erros de retorno do PR;",
            "Demais não integradas elegíveis (quando aplicável);",
            "Integradas recentes (quando houver conteúdo e regra de envio único).",
        ]
    )
    doc.add_paragraph(
        "Pendências de de-para e lote podem se repetir a cada ciclo até a correção. "
        "Alguns tipos de ocorrência entram uma única vez no histórico de envio para evitar spam."
    )

    # 10
    doc.add_heading("10. Perfis de acesso", level=1)
    doc.add_heading("10.1 Administrador", level=2)
    add_bullets(
        [
            "Visão de todos os estabelecimentos;",
            "Gestão de usuários, configurações, logs e auditoria;",
            "Pode forçar envio de relatório e alterar intervalo do e-mail.",
        ]
    )
    doc.add_heading("10.2 Usuário de estabelecimento", level=2)
    add_bullets(
        [
            "Opera apenas a própria unidade;",
            "Emite, acompanha, reemite e gerencia destinatários da unidade;",
            "Acessa dashboard e ajuda;",
            "Não acessa configurações globais nem cadastro amplo de usuários.",
        ]
    )

    # 11
    doc.add_heading("11. Ressalvas e pontos de atenção", level=1)
    add_bullets(
        [
            "Cadastro de de-para no PR é pré-requisito para materiais;",
            "Itens que exigem lote precisam estar corretos no Tasy;",
            "Local de estoque 104 não entra na integração;",
            "Nota já integrada no Tasy não volta à captura automática;",
            "Divergência de valores (total NF × soma dos itens), inclusive por desconto, pode ser rejeitada pelo PR;",
            "O painel não é espelho completo do Tasy: mostra o que passou pelo integrador;",
            "Acesso deve ser individual — há auditoria de ações e IP;",
            "Utilizar sempre a URL oficial do painel fornecida pela operação.",
        ]
    )

    # 12
    doc.add_heading("12. Glossário", level=1)
    add_table(
        ["Termo", "Significado"],
        [
            ["Tasy", "Sistema de origem das notas fiscais"],
            ["PR", "Sistema de destino (estoque/materiais)"],
            ["NR Sequência", "Identificador sequencial da nota no Tasy"],
            ["De-para", "Vínculo entre código de material do Tasy e do PR"],
            [
                "dt_integracao",
                "Data/hora em que a nota foi marcada como integrada no Tasy",
            ],
            ["Emitir pendentes", "Captura em lote das notas elegíveis da unidade"],
            ["Emitir nota específica", "Captura pontual por NR Sequência"],
            ["Reemitir", "Nova tentativa manual após falha"],
            ["Scheduler", "Rotina automática de captura (padrão 6 minutos)"],
            ["Sent", "Status de sucesso da integração no PR"],
        ],
    )

    doc.add_paragraph()
    fim = doc.add_paragraph()
    fim.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fim.add_run(
        "— Fim do documento —\n"
        "Documento elaborado pela GHR Tech para o Instituto Mais Saúde.\n"
        "Para dúvidas operacionais, utilize a aba Ajuda no painel ou o suporte da integração."
    )
    run.italic = True
    run.font.size = Pt(10)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)
    logo_ghr = _resolve_logo_ghr()
    if logo_ghr is None:
        print("AVISO: logo GHR nao encontrada em docs/assets/")
    else:
        print("Logo GHR:", logo_ghr.name)


if __name__ == "__main__":
    main()
